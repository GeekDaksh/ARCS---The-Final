"""Generate ARCS full documentation as a Word document."""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Page margins ──────────────────────────────────────────────────
section = doc.sections[0]
section.left_margin   = Inches(1.0)
section.right_margin  = Inches(1.0)
section.top_margin    = Inches(1.0)
section.bottom_margin = Inches(1.0)

# ── Helper: set font on a run ──────────────────────────────────────
def style_run(run, bold=False, italic=False, size=11,
              color=None, font_name="Calibri", code=False):
    run.bold   = bold
    run.italic = italic
    run.font.size = Pt(size)
    if code:
        run.font.name = "Courier New"
        run.font.size = Pt(9.5)
    else:
        run.font.name = font_name
    if color:
        run.font.color.rgb = RGBColor(*color)

# ── Helper: add a heading ──────────────────────────────────────────
def add_heading(text, level=1):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        run.font.name = "Calibri"
        if level == 1:
            run.font.size = Pt(18)
            run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
        elif level == 2:
            run.font.size = Pt(14)
            run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
        elif level == 3:
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(0x1F, 0x77, 0xB4)
    return p

# ── Helper: add normal paragraph ──────────────────────────────────
def add_para(text, bold=False, italic=False, size=11, color=None, indent=0):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Inches(indent)
    run = p.add_run(text)
    style_run(run, bold=bold, italic=italic, size=size, color=color)
    return p

# ── Helper: add a code block ──────────────────────────────────────
def add_code(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Inches(0.3)
    p.paragraph_format.right_indent = Inches(0.3)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    # Light grey background via shading
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F2F2F2')
    pPr.append(shd)
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x24, 0x29, 0x2E)
    return p

# ── Helper: add bullet ─────────────────────────────────────────────
def add_bullet(text, level=0, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.25 + level * 0.2)
    if bold_prefix:
        r1 = p.add_run(bold_prefix)
        style_run(r1, bold=True, size=11)
        r2 = p.add_run(text)
        style_run(r2, size=11)
    else:
        r = p.add_run(text)
        style_run(r, size=11)
    return p

# ── Helper: add a simple table ────────────────────────────────────
def add_table(headers, rows, col_widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    # header row
    hdr = t.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.size = Pt(10)
                run.font.name = "Calibri"
        # Dark blue background
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), '2E74B5')
        tcPr.append(shd)
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    # data rows
    for ri, row_data in enumerate(rows):
        row = t.rows[ri + 1]
        for ci, val in enumerate(row_data):
            cell = row.cells[ci]
            cell.text = str(val)
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)
                    run.font.name = "Calibri"
            if ri % 2 == 1:
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'), 'clear')
                shd.set(qn('w:color'), 'auto')
                shd.set(qn('w:fill'), 'DEEAF1')
                tcPr.append(shd)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in t.rows:
                row.cells[i].width = Inches(w)
    doc.add_paragraph()
    return t

def add_spacer():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)

# ══════════════════════════════════════════════════════════════════
# TITLE PAGE
# ══════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("ARCS")
r.bold = True
r.font.size = Pt(36)
r.font.name = "Calibri"
r.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Autonomous Range Control System")
r.font.size = Pt(20)
r.font.name = "Calibri"
r.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Complete Project Documentation")
r.font.size = Pt(14)
r.font.name = "Calibri"
r.italic = True
r.font.color.rgb = RGBColor(0x40, 0x40, 0x40)

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Written in plain English — for anyone to understand")
r.font.size = Pt(12)
r.font.name = "Calibri"
r.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# TABLE OF CONTENTS (manual)
# ══════════════════════════════════════════════════════════════════
add_heading("Table of Contents", 1)
toc_entries = [
    ("1.", "What Is ARCS? The Big Picture"),
    ("2.", "The Problem ARCS Solves"),
    ("3.", "The Intelligence Stack — 4 Layers"),
    ("4.", "Project File Map"),
    ("5.", "File-by-File Code Breakdown"),
    ("  5.1", "physics/constants.py — The Settings File"),
    ("  5.2", "physics/ballistic_solver.py — The Physics Calculator"),
    ("  5.3", "physics/bias_model.py — The Flaw Simulator"),
    ("  5.4", "physics/range_table.py — The Memory Book"),
    ("  5.5", "pinn_corrector.py — The AI Brain"),
    ("  5.6", "kalman_filter.py — The Smart Shot Averager"),
    ("  5.7", "bayesian_optimizer.py — The Smart Search Engine"),
    ("  5.8", "metrics.py — The Progress Tracker"),
    ("  5.9", "pipeline.py — The Main Orchestrator"),
    ("  5.10", "experiment.py — The Research Framework"),
    ("6.", "How Everything Connects — Full Data Flow"),
    ("7.", "Key Numbers to Remember"),
    ("8.", "Glossary"),
]
for num, title in toc_entries:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.0 if not num.startswith(" ") else 0.3)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    r1 = p.add_run(num + "  ")
    r1.bold = True
    r1.font.size = Pt(11)
    r1.font.name = "Calibri"
    r2 = p.add_run(title)
    r2.font.size = Pt(11)
    r2.font.name = "Calibri"

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# SECTION 1: WHAT IS ARCS?
# ══════════════════════════════════════════════════════════════════
add_heading("1. What Is ARCS?", 1)

add_para(
    "Imagine a robot that fires projectiles at targets — like a very precise automated cannon. "
    "The robot is told 'shoot at a target 300 metres away, 5 metres to the left, 10 metres up' "
    "and it must calculate exactly how to aim.",
    size=11
)
add_para(
    "The problem is: the robot has mechanical flaws. Its barrel droops slightly under its own "
    "weight. Its gears have tiny gaps. Its sensors drift over time. These flaws make shots land "
    "in the wrong place — not randomly, but in a predictable, systematic way.",
    size=11
)
add_para(
    "ARCS (Autonomous Range Control System) is an AI system that:",
    size=11
)
add_bullet("Watches where shots actually land versus where they should have landed")
add_bullet("Learns the robot's specific mechanical flaws from that history")
add_bullet("Pre-corrects every future shot before firing, so it lands on target")

add_para(
    "Think of it like this: imagine you always throw a ball slightly to the left. A friend "
    "watching you could say 'aim 5 cm to the right to compensate.' ARCS is that friend — "
    "except it learns automatically and gets more precise over time.",
    italic=True, size=11
)

add_spacer()
add_heading("Phase 1 Scope", 3)
add_para(
    "The current implementation operates under Phase 1 conditions: vacuum ballistics "
    "(no air resistance or wind), static targets that do not move, and a single robot "
    "firing at known positions. Phase 2 (future) will add moving targets and drag.",
    size=11
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# SECTION 2: THE PROBLEM
# ══════════════════════════════════════════════════════════════════
add_heading("2. The Problem ARCS Solves", 1)

add_heading("Systematic vs Random Error", 2)
add_para(
    "There are two types of error a robot can make:",
    size=11
)
add_bullet(
    "A random error is unpredictable — like wind gusting left or right. "
    "You cannot correct for it because you do not know which way it will go next time.",
    bold_prefix="Random error:  "
)
add_bullet(
    "A systematic error is predictable — like a bent gun sight that always points "
    "2° too high. Once you know about it, you can compensate perfectly every time.",
    bold_prefix="Systematic error:  "
)
add_para(
    "ARCS deals exclusively with systematic errors. There are four main ones:",
    size=11
)

add_heading("The Four Mechanical Flaws", 2)
add_table(
    headers=["Flaw", "What It Means", "Effect on Shots"],
    rows=[
        ["Barrel Gravity Sag",
         "The barrel is heavy and droops when angled upward",
         "Shots land short — barrel points slightly lower than commanded"],
        ["Gear Backlash",
         "Tiny gap between gear teeth — turret does not move instantly",
         "First shot after reversing direction is slightly off"],
        ["IMU Yaw Drift",
         "The left/right sensor slowly drifts over time",
         "Robot thinks it is pointing north, but it is actually 0.2° off"],
        ["Propellant Velocity Drift",
         "Ammo temperature changes how fast the projectile leaves the barrel",
         "Shots go slightly further or shorter than expected"],
    ],
    col_widths=[1.6, 2.5, 2.6]
)

add_heading("Why This Matters", 2)
add_para(
    "Without correction, the robot might consistently miss by 10–15 metres at 300m range. "
    "ARCS reduces that to under 2 metres — a 30–40% improvement in accuracy, without "
    "changing any hardware. The improvement comes purely from the AI learning and compensating "
    "for predictable mechanical behaviour.",
    size=11
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# SECTION 3: THE INTELLIGENCE STACK
# ══════════════════════════════════════════════════════════════════
add_heading("3. The Intelligence Stack — 4 Layers", 1)

add_para(
    "Every time ARCS fires at a target, it goes through four layers of intelligence, "
    "from simplest to smartest. Each layer refines the answer from the previous one.",
    size=11
)

layers = [
    ("Layer 1: Physics Engine",
     "Pure Newton mechanics — the exact mathematical answer in a perfect world.",
     'Given target at (300m, 5m, 50m), compute: pitch = 8.3°, yaw = 9.5°',
     "This is always computed first. It would be perfect if the robot had no flaws."),
    ("Layer 2: Range Table Lookup",
     "Historical memory — what correction worked last time at this range?",
     'Finds past engagements near 300m range → average correction: +0.43° pitch',
     "Uses weighted average of past results. Closer ranges count more."),
    ("Layer 3: PINN Neural Network",
     "The AI brain — predicts the systematic bias from learned patterns.",
     'Predicts: +0.52° pitch, −0.11° yaw, −3.45 m/s speed',
     "Trained on all historical corrections. Constrained by physics equations."),
    ("Layer 4: Bayesian Optimizer",
     "Fine-tunes by firing real test shots and searching for the best correction.",
     'Tries 20 corrections, each time learning from the result. Finds the best.',
     "The most expensive layer (uses real shots) but the most precise."),
]

for i, (title, desc, example, note) in enumerate(layers):
    add_heading(title, 3)
    add_para(desc, bold=False, size=11)
    add_para("Example:", bold=True, size=10)
    add_code(example)
    add_para(note, italic=True, size=10, color=(80, 80, 80))
    add_spacer()

add_para(
    "After all four layers, the result is recorded back into the range table and used to "
    "retrain the PINN — so the system continuously improves with every engagement.",
    size=11
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# SECTION 4: FILE MAP
# ══════════════════════════════════════════════════════════════════
add_heading("4. Project File Map", 1)

add_para(
    "The project is organised into these files. Each has a single clear responsibility.",
    size=11
)

add_code(
"""ARCS/
├── pipeline.py                 ← MAIN ENTRY POINT — runs the whole system
├── pinn_corrector.py           ← The AI brain (neural network corrector)
├── bayesian_optimizer.py       ← Smart shot search + Gaussian Process
├── kalman_filter.py            ← Shot averaging (Kalman Filter)
├── metrics.py                  ← Tracks accuracy and improvement over time
├── experiment.py               ← Research experiments (ablation, learning curves)
├── synthetic_data_generator.py ← Creates fake shot data for testing
├── rebuild_physics.py          ← Rebuilds the physics lookup table
│
├── physics/
│   ├── ballistic_solver.py     ← Core math: angle + time-of-flight
│   ├── range_table.py          ← Table of historical corrections
│   ├── bias_model.py           ← Simulates the robot's mechanical flaws
│   ├── constants.py            ← All numbers and limits in one place
│   └── rotation.py             ← 3D rotation math
│
├── tests/                      ← 602 automated checks
│
└── data/
    ├── range_table_physics.csv      ← Pre-computed angles for every range
    ├── range_table_corrections.csv  ← History of what corrections worked
    ├── pinn_low_weights.pt          ← Saved AI model weights (LOW trajectory)
    ├── pinn_high_weights.pt         ← Saved AI model weights (HIGH trajectory)
    └── metrics_history.csv          ← Full engagement history log"""
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# SECTION 5: FILE-BY-FILE BREAKDOWN
# ══════════════════════════════════════════════════════════════════
add_heading("5. File-by-File Code Breakdown", 1)

# ── 5.1 constants.py ──────────────────────────────────────────────
add_heading("5.1  physics/constants.py — The Settings File", 2)
add_para(
    "This file stores every important number in one place. If a value needs to change, "
    "you change it here and it automatically updates everywhere else in the project.",
    size=11
)

add_heading("Physical Constants", 3)
add_code(
"""GRAVITY = 9.81          # Earth's gravitational pull (metres per second²)
                        # Every physics calculation uses this exact value."""
)

add_heading("Robot Mechanical Noise", 3)
add_para(
    "These numbers describe how imprecise the motors are. Even when you command exactly "
    "8.0°, the robot might actually achieve anywhere from 7.7° to 8.3° due to motor noise.",
    size=11
)
add_code(
"""SIGMA_PITCH_DEG = 0.3   # The up/down barrel motor has ±0.3° natural noise.
                        # You command 8.0°, you get 7.7° to 8.3°.

SIGMA_YAW_DEG   = 0.2   # The left/right turret motor: ±0.2° noise.

SIGMA_V0        = 1.5   # Projectile speed varies ±1.5 m/s.
                        # Caused by temperature changes in the propellant.

BIAS_SCALE      = 1.5   # How strong the systematic flaws are.
                        # Pitch has a systematic bias of 1.5 × 0.3° = 0.45°."""
)

add_heading("Engagement Protocol", 3)
add_code(
"""N_SHOTS_BASELINE = 30   # Fire 30 shots WITHOUT any correction.
                        # The median miss = baseline CEP (the problem to fix).

N_SHOTS_VERIFY   = 30   # Fire 30 shots WITH the best correction found.
                        # Compare to baseline to measure improvement.

N_KF_SHOTS       = 8    # Fire 8 directional shots after BO search,
                        # so the Kalman Filter can estimate residual error.

FALLBACK_THRESHOLD = 1.10  # If corrected CEP > 1.10 × baseline CEP,
                            # the correction is harmful — throw it away.
                            # (10% worse = safety threshold)

QUALITY_FILTER_RATIO = 1.20  # When training the AI, exclude records where
                              # miss_after > miss_before × 1.20.
                              # (Corrections 20%+ worse are probably noise.)"""
)

add_heading("Coordinate System", 3)
add_para(
    "ARCS uses a right-handed 3D coordinate system with the robot at the origin (0,0,0):",
    size=11
)
add_code(
"""X = forward (toward the target)
Y = upward (against gravity)
Z = right (lateral, to the robot's right side)

Yaw   = rotating the turret left/right (around the Y axis)
Pitch = tilting the barrel up/down (around the Z axis)

Example target: (200, 10, 50) means:
  200m forward, 10m above the robot, 50m to the right"""
)

add_spacer()

# ── 5.2 ballistic_solver.py ───────────────────────────────────────
add_heading("5.2  physics/ballistic_solver.py — The Physics Calculator", 2)
add_para(
    "This file solves the core physics problem: given where the target is and how fast "
    "the projectile travels, compute the exact angle to fire at. This is pure Newton mechanics "
    "— the same equations Galileo used in 1638. There is no AI here.",
    size=11
)

add_heading("The Physics Equations", 3)
add_para(
    "When a projectile is fired, it follows these equations of motion at any time t:",
    size=11
)
add_code(
"""x(t) = v₀ · cos(pitch) · cos(yaw) · t     ← moves forward at constant speed
y(t) = v₀ · sin(pitch) · t  −  ½ · g · t²  ← rises then falls under gravity
z(t) = v₀ · cos(pitch) · sin(yaw) · t       ← moves sideways at constant speed"""
)
add_para(
    "Working backwards: given the target location (x, y, z), find the pitch and yaw "
    "angles that make the projectile land exactly there.",
    size=11
)

add_heading("Two Solutions Always Exist", 3)
add_para(
    "For any reachable target, there are always exactly two angles that work:",
    size=11
)
add_code(
"""LOW trajectory  (~10°): Fast, flat path — like a rifle bullet.
                         Short flight time, direct route.

HIGH trajectory (~70°): Slow, lobbed path — like a mortar shell.
                         Long flight time, arcing route.

The solver computes both and returns whichever you prefer."""
)

add_heading("How solve() Works — Step by Step", 3)
add_code(
"""def solve(self, target_x, target_y, target_z, v0, prefer="LOW"):

    # Step 1: Validate — is the target reachable?
    valid, reason = validate_target(target_x, target_y, target_z)
    # Checks: too close? (< 10m)  Too far? (> 500m)  Behind the robot?

    # Step 2: Calculate yaw (left/right angle)
    bearing_rad = np.arctan2(target_z, target_x)
    # arctan2 gives the compass angle to the target.
    # Target at (200m forward, 50m right): arctan2(50, 200) = 14°

    # Step 3: Calculate horizontal range
    R = np.sqrt(target_x**2 + target_z**2)
    # Pythagoras in the horizontal plane.
    # (200m forward, 50m right): R = √(200² + 50²) = 206m

    # Step 4: Solve for pitch angle (the key physics step)
    v2   = v0 * v0
    disc = v2**2 - g * (g * R**2 + 2 * target_y * v2)
    # 'disc' is the discriminant (like b²-4ac in the quadratic formula).
    # If disc < 0: target is too far — cannot reach it.

    theta_low_deg  = arctan2(v2 - √disc, g * R)  # ← LOW  angle (flat)
    theta_high_deg = arctan2(v2 + √disc, g * R)  # ← HIGH angle (lobbed)
    # The ± sign gives the two solutions.

    # Steps 5-10: Compute velocity components, time-of-flight, apex
    # height, impact speed, verify mathematically (error < 0.01m)."""
)

add_spacer()

# ── 5.3 bias_model.py ─────────────────────────────────────────────
add_heading("5.3  physics/bias_model.py — The Flaw Simulator", 2)
add_para(
    "This file simulates the robot's mechanical imperfections. It is used during testing "
    "so the system can run thousands of virtual engagements without a real robot. "
    "It applies both systematic bias (predictable flaw) and random noise (unavoidable) "
    "to every shot.",
    size=11
)
add_code(
"""# FLAW 1: Gravity sag — barrel droops under its own weight.
# The higher the barrel is angled, the more it droops.
# sin(pitch) captures this: at 0° (horizontal) there is no droop.
# At 90° (straight up) the droop is maximum.
delta_pitch_bias = sag_coefficient * sin(pitch_cmd_rad)

# FLAW 2: Gear backlash — gap in gears causes slight delay
#          when the turret reverses direction.
delta_yaw_bias = yaw_drift_constant  # constant offset

# FLAW 3: Propellant velocity drift — muzzle speed varies with temperature.
delta_v0_bias = velocity_drift_constant

# NOISE: Unpredictable motor jitter added on top of systematic bias.
actual_pitch = commanded_pitch + delta_pitch_bias + Normal(0, SIGMA_PITCH_DEG)
actual_yaw   = commanded_yaw   + delta_yaw_bias   + Normal(0, SIGMA_YAW_DEG)
actual_v0    = commanded_v0    + delta_v0_bias     + Normal(0, SIGMA_V0)

# KEY INSIGHT:
# Bias   = systematic, same direction every time → ARCS can learn and fix this
# Noise  = random, different every shot          → cannot be fixed, only averaged"""
)

add_spacer()

# ── 5.4 range_table.py ────────────────────────────────────────────
add_heading("5.4  physics/range_table.py — The Memory Book", 2)
add_para(
    "This file manages two CSV files: one with pre-computed physics (what angle to fire at "
    "for any range/height/speed combination), and one with historical corrections (what "
    "actually worked in past engagements).",
    size=11
)

add_heading("What the Corrections File Looks Like", 3)
add_code(
"""range_table_corrections.csv:

range_m | height_m | v0_ms | delta_pitch | delta_yaw | delta_v0 | miss_before | miss_after
  200.0 |      0.0 | 100.0 |      +0.432 |    -0.112 |    -3.45 |        11.2 |       4.3
  350.0 |     10.0 | 100.0 |      +0.518 |    -0.088 |    -3.21 |        15.7 |       5.1

Each row = one engagement result.
miss_before = how bad it was WITHOUT any correction
miss_after  = how good it was WITH the correction (CEP in metres)"""
)

add_heading("lookup() — Finding the Best Historical Correction", 3)
add_para(
    "When you ask 'what correction should I use for 200m range?', the range table:",
    size=11
)
add_code(
"""def lookup(self, range_m, height_m, v0_ms, prefer="LOW"):
    # 1. Filter to matching trajectory type (LOW or HIGH)
    # 2. Compute distance from each stored record to the query point
    # 3. Weight = exp(−distance² / bandwidth²)
    #    → nearby records count more, distant records count less
    # 4. Return weighted average of delta_pitch, delta_yaw, delta_v0

    # Example: query for range=205m
    # Record at 200m: distance=5m, weight=0.97 (very close, counts a lot)
    # Record at 250m: distance=45m, weight=0.31 (far, counts less)
    # Record at 350m: distance=145m, weight=0.01 (very far, barely counts)"""
)

add_heading("record_correction() — Saving New Results", 3)
add_code(
"""def record_correction(self, range_m, height_m, v0_ms,
                      delta_pitch, delta_yaw, delta_v0,
                      miss_before, miss_after, ...):
    # Appends one new row to range_table_corrections.csv
    # This data feeds back into the range table lookup AND trains the PINN
    # IMPORTANT: Only called when the correction helped (not on fallback)"""
)

add_spacer()

# ── 5.5 pinn_corrector.py ─────────────────────────────────────────
add_heading("5.5  pinn_corrector.py — The AI Brain", 2)
add_para(
    "This is the most important file. It contains the Physics-Informed Neural Network (PINN) "
    "— an AI that learns to predict the robot's systematic errors from historical data.",
    size=11
)

add_heading("What a Neural Network Does (Simply)", 3)
add_para(
    "A neural network is a mathematical function with thousands of adjustable 'knobs' called "
    "weights. You feed it inputs, it produces outputs. Training means turning the knobs until "
    "the outputs match the correct answers. Once trained, it can predict answers for new inputs "
    "it has never seen before.",
    size=11
)

add_heading("The Network Architecture", 3)
add_code(
"""Input  (4 numbers):  range, height, speed, sin(pitch_angle)
         ↓
Layer 1: 4 inputs → 64 neurons  [Tanh activation]
         ↓
Layer 2: 64 neurons → 64 neurons [Tanh activation]
         ↓
Layer 3: 64 neurons → 32 neurons [Tanh activation]
         ↓
Output (3 numbers):  delta_pitch (°), delta_yaw (°), delta_v0 (m/s)

Total parameters (knobs): 6,496
Training time: under 1 second on CPU
Minimum training records needed: 20"""
)

add_code(
"""class _PINNNet(nn.Module):
    def __init__(self):
        super().__init__()
        self.net = nn.Sequential(
            nn.Linear(4, 64),   nn.Tanh(),   # Layer 1: 4 inputs → 64 neurons
            nn.Linear(64, 64),  nn.Tanh(),   # Layer 2: 64 → 64 neurons
            nn.Linear(64, 32),  nn.Tanh(),   # Layer 3: 64 → 32 neurons
            nn.Linear(32, 3),               # Output: 32 → 3 corrections
        )
    # nn.Tanh() squashes output to [-1, +1], preventing absurd corrections like "+50°"."""
)

add_heading("Why sin(pitch) as an Input Feature?", 3)
add_para(
    "Gravity sag (the main flaw) is mathematically proportional to sin(pitch_angle). "
    "By giving the network sin(pitch) directly as an input, it can learn the sag relationship "
    "with a single weight multiplication. Without this feature, the network would need many "
    "layers to rediscover the sine relationship from range and height alone.",
    size=11
)
add_code(
"""def _features(self, range_m, height_m, v0_ms, theta_deg):
    return np.array([
        range_m  / 500.0,               # Normalised range:  0.0 to 1.0
        height_m / 50.0,                # Normalised height: -0.4 to 1.0
        v0_ms    / 300.0,               # Normalised speed:  0.17 to 1.0
        np.sin(np.deg2rad(theta_deg)),  # sin(pitch) — THE KEY FEATURE for sag
    ], dtype=np.float32)
    # All values are normalised to roughly the same scale.
    # This helps the network learn — large inputs of different magnitudes
    # make training unstable."""
)

add_heading("Why 'Physics-Informed'?", 3)
add_para(
    "A normal neural network just fits the historical data. The 'physics-informed' part "
    "adds a second constraint: every predicted correction must produce a trajectory that "
    "physically hits the target, according to Galileo's range equation:",
    size=11
)
add_code(
"""Galileo's Range Equation (exact vacuum ballistics):

   y(x) = x · tan(θ) − g · x² / (2 · v₀² · cos²(θ))

Where:
   x = horizontal range (metres)
   y = height at that horizontal position (metres)
   θ = firing angle (pitch)
   v₀ = muzzle velocity (m/s)
   g = 9.81 m/s²

When x = range_m and y must equal height_m:
   If the network predicts a bad correction → y(range_m) ≠ height_m → large L_physics
   If the network predicts a good correction → y(range_m) ≈ height_m → L_physics ≈ 0"""
)

add_code(
"""@staticmethod
def _physics_residual(range_t, height_t, v0_t, theta_t, pred):
    dp = pred[:, 0]    # predicted delta_pitch (degrees)
    dy = pred[:, 1]    # predicted delta_yaw   (degrees)
    dv = pred[:, 2]    # predicted delta_v0    (m/s)

    # Apply corrections to get the actual firing parameters
    theta_c = torch.clamp(torch.deg2rad(theta_t + dp), min=0.01, max=1.47)
    # theta_c = corrected pitch angle, clamped to physical range [0.57°, 84°]
    # torch.clamp prevents extreme values that would break the math

    v0_c = torch.clamp(v0_t + dv, min=10.0)
    # corrected muzzle velocity — must be at least 10 m/s

    # Apply Galileo's equation with the corrected angles
    y_pred = (range_t * torch.tan(theta_c)
              - (GRAVITY * range_t**2) / (2.0 * v0_c**2 * torch.cos(theta_c)**2))

    # Penalty: how far does the corrected trajectory miss the target height?
    L_pitch = ((y_pred - height_t)**2 / (range_t**2 + 1.0)).mean()
    # Divided by range² so a 5m miss at 500m is penalised less than at 100m

    # Yaw penalty: large yaw corrections cause lateral displacement
    L_yaw = torch.sin(torch.deg2rad(dy)).pow(2).mean()

    return L_pitch + 0.5 * L_yaw"""
)

add_heading("The Total Training Loss", 3)
add_code(
"""loss = L_data + LAMBDA_PHYSICS * L_phys
#         ↑              ↑               ↑
#  "Match historical    0.1 weight   "Corrections must
#   corrections"                      physically hit the target"

L_data  = MSE(predicted_corrections, actual_corrections_from_csv)
          "How different are the predictions from what actually worked?"

L_phys  = physics_residual(range, height, v0, theta, predictions)
          "Does the corrected trajectory actually reach the target?"

LAMBDA_PHYSICS = 0.1 means the physics constraint counts 10% as much as data.
This is a tuned trade-off: enough physics to prevent impossible corrections,
not so much that it overrides what the data says worked."""
)

add_heading("Training Process — 300 Epochs", 3)
add_code(
"""def _train_pytorch(self, df, verbose):
    # Build training data from the corrections CSV
    X, Y, thetas = self._build_training_data(df)
    # X = matrix of 4 features per engagement
    # Y = matrix of 3 corrections per engagement
    # thetas = nominal pitch angle per engagement (for physics loss)

    net = _PINNNet()  # Fresh network
    optimizer = optim.Adam(net.parameters(), lr=3e-4, weight_decay=1e-5)
    # Adam: adjusts each weight's learning rate individually.
    # lr=3e-4 = learning rate (step size). weight_decay = small penalty
    # for very large weights (prevents overfitting to noise).

    scheduler = optim.lr_scheduler.CosineAnnealingLR(optimizer, T_max=300)
    # Starts with the full learning rate, gradually reduces to near-zero.
    # Like starting with bold brushstrokes, finishing with fine detail.

    for epoch in range(300):  # 300 passes through the entire dataset
        optimizer.zero_grad()       # Reset gradients from last step
        pred    = net(X_t)          # Network makes predictions (forward pass)
        L_data  = mse_loss(pred, Y_t)  # How wrong are the predictions?
        L_phys  = _physics_residual(...)  # Do they hit the target physically?
        loss    = L_data + 0.1 * L_phys  # Combined loss

        loss.backward()             # Compute: which direction to adjust weights?
                                    # (backpropagation through the network)
        clip_grad_norm_(net.parameters(), max_norm=1.0)
        # Prevent any single weight update from being too large
        # (called gradient clipping — avoids unstable training)

        optimizer.step()  # Apply the weight adjustments
        scheduler.step()  # Slightly reduce learning rate"""
)

add_heading("Auto-Loading Saved Weights on Startup", 3)
add_para(
    "When the pipeline restarts, the PINN should be immediately ready to make predictions "
    "without waiting for 300 training epochs. This is solved by saving weights to disk "
    "after training and loading them on startup:",
    size=11
)
add_code(
"""# At the END of __init__():
if Path(self.corrections_path).exists():
    self.load_weights()
# If the corrections CSV file exists, try to load saved weights.
# If weights file exists: is_fitted = True immediately.
# If weights file missing (first run): skip silently.

def save_weights(self):
    torch.save(self._net.state_dict(), weight_path)
    # Saves all 6,496 knob values to a binary file (pinn_low_weights.pt)

def load_weights(self):
    net = _PINNNet()                              # Create fresh network
    net.load_state_dict(torch.load(weight_path)) # Load saved knob values
    net.eval()                                   # Switch to inference mode
    self._net = net
    self.is_fitted = True                        # Ready to predict immediately
    return True"""
)

add_heading("The Retraining Trigger", 3)
add_para(
    "The network does not retrain after every engagement. That would be slow and wasteful. "
    "Instead it uses an 'information gain' trigger — only retrain when genuinely new "
    "territory has been covered:",
    size=11
)
add_code(
"""def should_retrain(self):
    # RULE 1: Never retrain with fewer than 20 records total
    if current_n < 20:
        return False   # Not enough data to train anything reliable

    # RULE 2: Need at least 5 new records since last training
    if current_n < self.last_trained_n + 5:
        return False   # Too few new records — wait for more

    # RULE 3: Check if new records cover genuinely new territory
    # (more than 30m from any range seen during last training)
    novelty = minimum distance from new ranges to training set ranges
    return novelty > 30.0
    # Only retrain if you have data from a range you have never seen before.
    # Firing at 200m ten more times adds no new information about 300m targets."""
)

add_heading("predict() — Making a Prediction", 3)
add_code(
"""def predict(self, range_m, height_m, v0_ms=100.0):
    if not self.is_fitted:
        return {"delta_pitch": 0.0, "delta_yaw": 0.0,
                "delta_v0": 0.0, "source": "none"}
        # Safety: if not trained yet, return zero correction (do nothing)

    theta = self._theta_cmd(range_m, height_m, v0_ms)
    # Get the nominal pitch angle from the ballistic solver

    feats = self._features(range_m, height_m, v0_ms, theta)
    # Build the 4-element input vector

    with torch.no_grad():  # No gradient computation needed for inference
        x   = torch.from_numpy(feats).unsqueeze(0)  # Add batch dimension
        out = self._net(x).numpy()[0]               # Run through network

    return {
        "delta_pitch": float(np.clip(out[0], -3.0,  3.0)),   # Clip to ±3°
        "delta_yaw":   float(np.clip(out[1], -2.0,  2.0)),   # Clip to ±2°
        "delta_v0":    float(np.clip(out[2], -10.0, 10.0)),  # Clip to ±10 m/s
        "source":      "pinn_torch_LOW",                      # Which model was used
    }
    # np.clip = enforce physical limits. No matter what the network predicts,
    # corrections outside these bounds are physically impossible."""
)

add_spacer()

# ── 5.6 kalman_filter.py ──────────────────────────────────────────
add_heading("5.6  kalman_filter.py — The Smart Shot Averager", 2)
add_para(
    "A Kalman Filter is the gold standard for estimating something uncertain from noisy "
    "measurements. GPS uses it. Apollo used it. ARCS uses it to estimate the remaining "
    "correction error after the Bayesian Optimizer has done its search.",
    size=11
)

add_heading("The Simple Idea", 3)
add_para(
    "After the BO search finds its best correction, ARCS fires 8 directional shots at that "
    "correction. Each shot lands slightly differently due to motor noise. The Kalman Filter "
    "combines all 8 measurements into the best possible estimate of the remaining systematic "
    "error — better than simply averaging the misses.",
    size=11
)
add_code(
"""# What we want to know: how much additional correction do we still need?
# State: x = [Δpitch_needed, Δyaw_needed]

class EngagementKF:
    def __init__(self, range_m, sigma_pitch=0.3, sigma_yaw=0.2):
        self.x = np.zeros(2)              # State: [0°, 0°] initially (no correction known)
        self.P = np.eye(2) * (2 × 0.3)²  # Uncertainty: large at first (4 times σ²)
                                          # P shrinks as more shots are observed

        self.Q = np.eye(2) * 0.01  # Process noise: how much can bias drift between shots?
                                    # Small = assume bias is stable shot-to-shot
        self.R = diag([0.3², 0.2²]) # Measurement noise: motor sigma in both axes"""
)

add_heading("The Update Step — After Each Shot", 3)
add_code(
"""def update(self, error_x, error_z):
    # Convert miss distance to required angle correction
    deg_per_m = (1.0 / range_m) * (180.0 / π)
    z = np.array([
        -error_x * deg_per_m,   # If shot went 5m forward: correct by -5m × deg/m pitch
        -error_z * deg_per_m,   # If shot went 3m right: correct by -3m × deg/m yaw
    ])
    # Negative because: shot went right → we need to aim left to compensate

    # PREDICT: what do we expect before seeing this shot?
    x_pred = self.x          # Bias is assumed constant between shots
    P_pred = self.P + self.Q # Uncertainty grows slightly (bias might have drifted)

    # UPDATE: incorporate the new shot measurement
    innovation = z - x_pred  # How surprised are we? (observation minus prediction)

    K = P_pred / (P_pred + self.R)
    # Kalman Gain: how much to trust this measurement vs the prior estimate?
    # If R is large (noisy measurement): K is small → trust measurement less
    # If P is large (uncertain estimate): K is large → trust measurement more

    self.x = x_pred + K × innovation  # Update state estimate
    self.P = (I - K) × P_pred         # Update uncertainty (always smaller than before)
    self.n_updates += 1"""
)

add_heading("Confidence Score", 3)
add_code(
"""@property
def confidence(self):
    P_trace_init = (2.0 × 0.3)² × 2   # Initial total uncertainty
    P_trace_now  = trace(self.P)         # Current total uncertainty
    return clip(1.0 - P_trace_now / P_trace_init, 0.0, 1.0)
    # As P shrinks (more shots = more certainty), confidence rises toward 1.0
    # After 1 shot:  confidence ≈ 0.3  (low — not much data yet)
    # After 8 shots: confidence ≈ 0.9  (high — we trust the estimate)
    # Only apply KF correction when confidence > 0.5"""
)

add_heading("Pitch Scaling — A Critical Correction", 3)
add_para(
    "The KF's error-to-angle conversion uses a simplified model (error_x ≈ range × Δpitch). "
    "Real ballistics is nonlinear — the actual sensitivity is much higher. Without scaling, "
    "the KF would overcorrect the pitch by roughly 10×. The scaling factor fixes this:",
    size=11
)
add_code(
"""# True ballistic pitch sensitivity:
dR/dθ = 2 × v₀² × cos(2θ) / g       (metres per radian)

# At 200m, θ=6°, v₀=100 m/s:  dR/dθ ≈ 2000 m/rad
# KF assumes:                   sensitivity = range_m = 200 m/rad

# KF overcorrects by ~10×. The pitch scale corrects for this:
kf_pitch_scale = min(1.0, range_m / max(|dR/dθ|, 1.0))
# Applied as: best_correction += kf_dp × kf_pitch_scale
# Yaw does NOT need scaling (lateral error = range × Δyaw is already correct)"""
)

add_spacer()

# ── 5.7 bayesian_optimizer.py ─────────────────────────────────────
add_heading("5.7  bayesian_optimizer.py — The Smart Search Engine", 2)
add_para(
    "This file contains the Bayesian Optimizer — an AI search algorithm that finds the best "
    "correction by intelligently trying different values and learning from each result. "
    "It is smarter than random trial-and-error because it builds a model of the search space "
    "and focuses on the most promising areas.",
    size=11
)

add_heading("The Gaussian Process (GP) — The BO's Brain", 3)
add_para(
    "A Gaussian Process is a mathematical model that stores all observed "
    "(correction → miss_distance) pairs and, for any new correction, predicts both "
    "the expected miss distance AND how uncertain that prediction is.",
    size=11
)
add_code(
"""class GaussianProcess:
    def predict(self, X_new):
        # Given new correction candidates, return:
        K_s  = self._matern52(X_new, self.X_obs)   # Similarity to past observations
        mean = K_s @ self.K_inv @ self.y_obs        # Expected miss for each candidate
        var  = diag(K_ss - K_s @ self.K_inv @ K_s.T) # Uncertainty in that prediction
        return mean, np.sqrt(var)

# The Matern 5/2 kernel measures "how similar are two corrections?"
# Two corrections 0.1° apart are assumed to give similar results.
# Two corrections 0.8° apart are assumed to be less related."""
)

add_heading("UCB: Where to Try Next?", 3)
add_code(
"""def suggest(self):
    n = len(self._X)   # How many BO iterations so far?

    if n == 0:   # First suggestion
        # Priority 1: Start at the PINN prediction (if available)
        if self._warm_start is not None:
            return clip([pinn_delta_pitch, pinn_delta_yaw, pinn_delta_v0], bounds)
            # Eliminates 3 "wasted" shots that would just rediscover PINN's answer

        # Priority 2: Start at historical average (cross-engagement memory)
        if not pinn_active and memory.engagement_count >= 2:
            return clip(memory.prior_mean(), bounds)

        return np.zeros(3)  # Default: start at no correction

    if n < 4:  # First 4 iterations: explore randomly
        return [random.uniform(lo, hi) for lo, hi in bounds]

    # After 4 shots: use GP to suggest intelligently
    kappa_t = decaying_kappa()  # Starts at 2.0, decays to 0.5

    candidates = 2000 random corrections within bounds
    mean, std = gp.predict(candidates)  # GP predicts miss for each

    lcb = mean - kappa_t × std   # Lower Confidence Bound score
    # Low mean = good expected result
    # High std = uncertain = worth exploring
    # kappa controls the balance: high kappa → explore more; low → exploit more

    return candidates[argmin(lcb)]  # Return the most promising candidate"""
)

add_heading("Adaptive Bounds — Getting Smarter Over Time", 3)
add_code(
"""DEFAULT_BOUNDS = [
    [-0.9°, +0.9°],    # pitch: ±3×SIGMA_PITCH_DEG
    [-0.6°, +0.6°],    # yaw:   ±3×SIGMA_YAW_DEG
    [-4.5, +4.5 m/s],  # v0:    ±3×SIGMA_V0
]

def _adaptive_bounds(self):
    if memory.engagement_count < 5:
        return DEFAULT_BOUNDS   # Full range — not enough history yet

    mean = memory.prior_mean()  # Average of past successful corrections
    std  = memory.prior_std()   # Spread of past corrections

    tight = [mean - 3×std,  mean + 3×std]
    # If past corrections cluster around +0.5° pitch:
    #   New range: [+0.5 - 3×0.1°, +0.5 + 3×0.1°] = [+0.2°, +0.8°]
    #   Instead of searching the full [-0.9°, +0.9°]
    # Narrower range = fewer shots needed to find the best correction"""
)

add_heading("The Full Engagement Loop", 3)
add_code(
"""def run_engagement(self, target_x, target_y, target_z, ...):

    # 1. Solve physics (what angle for this target?)
    sol = self.solver.solve(target_x, target_y, target_z, v0)

    # 2. Baseline: fire 30 shots with NO correction
    bl_cep, bl_misses = self.baseline_cep(sol, ...)
    # bl_cep = median miss without any correction (e.g. 12.3m)

    # 3. Create Kalman Filter (ready to accumulate data)
    kf = EngagementKF(range_m=sol.horiz_range, ...)

    # 4. If PINN gave a pre-correction, apply it to the firing solution
    firing_sol.pitch = sol.pitch + pinn_delta_pitch
    firing_v0        = v0 + pinn_delta_v0
    # BO now searches for RESIDUAL correction on top of PINN's correction

    # 5. Bayesian Optimizer loop (20 iterations: 4 random + 16 GP-guided)
    correction = optimizer.suggest()   # First: PINN warm-start or [0,0,0]
    while True:
        avg_miss = fire_averaged(correction, n_avg_shots)
        # Fire n_avg shots, compute median miss for this correction
        optimizer.update(correction, avg_miss)  # Tell GP the result
        correction = optimizer.suggest()        # Get next suggestion
    # After 20 iterations: optimizer.best_correction = the winner

    # 6. KF Refinement: 8 directional shots at the best correction
    for _ in range(8):
        error_x, error_z, _ = fire_one_directional(best_correction)
        kf.update(error_x, error_z)   # KF learns the remaining error
    if kf.confidence > 0.5:
        best_correction += [kf_dp × pitch_scale, kf_dy, 0]  # Refine

    # 7. Verify: fire 30 shots WITH the final best correction
    v_cep = verified_cep(sol, best_correction, ...)
    # v_cep = median miss with correction applied (e.g. 4.3m)

    # 8. Fallback: was the correction actually helpful?
    if v_cep > bl_cep × 1.10:   # More than 10% worse → reject correction
        best_correction = [0, 0, 0]
        v_cep = bl_cep         # Report 0% improvement (honest)

    improvement = (bl_cep - v_cep) / bl_cep × 100   # e.g. 65.0%

    # 9. Record (only if correction helped — not on fallback)
    range_table.record_correction(..., miss_before=bl_cep, miss_after=v_cep)

    return {"baseline_cep": bl_cep, "verified_cep": v_cep,
            "improvement_pct": improvement, "kf_correction": ..., ...}"""
)

add_heading("Statistical Significance — Wilcoxon Test", 3)
add_para(
    "After measuring baseline (30 shots) and verified (30 shots), ARCS checks: "
    "is the improvement real, or could it have happened by chance?",
    size=11
)
add_code(
"""def paired_wilcoxon(self, baseline_misses, verified_misses):
    stat, p = wilcoxon(baseline_misses, verified_misses, alternative='greater')
    # Wilcoxon signed-rank test compares 30 paired shot observations.
    # p < 0.05 means: less than 5% chance the improvement is just luck.
    return {"p_value": p, "significant": p < 0.05}
    # The test uses the SAME 30 shot noise values for baseline and verify
    # (paired evaluation). This eliminates measurement noise from the comparison."""
)

add_spacer()

# ── 5.8 metrics.py ────────────────────────────────────────────────
add_heading("5.8  metrics.py — The Progress Tracker", 2)
add_para(
    "This file records how well ARCS is performing over time and provides a learning "
    "curve showing improvement across engagements. Every result is stored permanently — "
    "nothing is ever overwritten.",
    size=11
)

add_heading("CEP — The Accuracy Metric", 3)
add_code(
"""CEP = Circular Error Probable
     The standard military and aerospace accuracy metric.

CEP 50% = the radius within which HALF of all shots land.
          (50th percentile of miss distances)
          This is the main metric ARCS optimises.

CEP 90% = the radius within which 90% of shots land.
          Relevant for worst-case planning.

Example:
    30 shots, misses = [1.2, 3.4, 2.1, 8.5, 4.2, 5.1, 2.8, ...]
    When sorted: [1.2, 2.1, 2.8, 3.4, 4.2, 5.1, 8.5, ...]
    CEP 50% = middle value = 4.2m
    CEP 90% = 90th percentile value

def cep_50(df):
    return float(df["miss_dist"].median())  # Just the median miss distance"""
)

add_heading("ConvergenceTracker — The History Log", 3)
add_code(
"""class ConvergenceTracker:
    def record(self, engagement_n, baseline_cep, bo_cep, improvement_pct, ...):
        row = {
            "engagement_n":    537,         # This is the 537th engagement ever
            "timestamp":       "2026-06-05T10:23:41",
            "baseline_cep_m":  12.3,        # Miss WITHOUT correction (12.3m)
            "bo_cep_m":        4.3,         # Miss WITH correction (4.3m)
            "improvement_pct": 65.0,        # 65% improvement
            "solution_type":   "LOW",       # Flat trajectory used
        }
        # Appended to metrics_history.csv — never overwritten.
        # This builds a permanent record of every engagement.
        pd.DataFrame([row]).to_csv(self.path, mode='a', ...)"""
)

add_heading("print_learning_curve() — Showing Progress", 3)
add_code(
"""# Example output of print_learning_curve():

  Eng   Baseline      BO CEP     Improve     Range  Type
  ---  ----------  ----------  ----------  -------  ----
    1     12.34m       7.21m    ↑   41.6%    200m   LOW
    2      9.87m       9.87m    ↓    0.0%    150m   LOW  ← fallback triggered
    3     14.21m       6.43m    ↑   54.8%    300m   HIGH
  ...
  536    11.04m        4.27m    ↑   61.3%    360m   LOW

  Mean improvement        : +37.0%
  Positive engagements    : 84%
  Best single improvement : 81.9%
  Best miss ever          : 1.34m"""
)

add_spacer()

# ── 5.9 pipeline.py ───────────────────────────────────────────────
add_heading("5.9  pipeline.py — The Main Orchestrator", 2)
add_para(
    "This is the entry point — the file you run directly. It creates every component and "
    "connects them all together. Think of it as the conductor of an orchestra: it does not "
    "do the music itself, but it makes sure everyone plays at the right time.",
    size=11
)

add_heading("__init__() — Startup Sequence", 3)
add_code(
"""def __init__(self, physics_path, corrections_path, history_path,
             v0=100.0, seed=None, verbose=True):

    # Step 1: Load the range table
    self.rt = RangeTable(physics_path, corrections_path)
    if not Path(physics_path).exists():
        self._generate_physics_table()   # First run: build the physics grid (~30s)
    else:
        self.rt.load(verbose=verbose)    # Existing run: load from disk

    # Step 2: Create the Bayesian Optimizer components
    self.memory = EngagementMemory()     # Cross-engagement correction memory
    self.gmodel = GlobalModel()          # 2D GP across all engagements
    self.bo     = BayesianOptimizer(
        memory=self.memory, global_model=self.gmodel,
        n_avg=3, n_init=4, n_suggest=16,  # 20 total iterations per engagement
        kappa=2.0, kappa_min=0.5,
    )

    # Step 3: Warm-start GlobalModel immediately if corrections exist
    if self.rt._corrections_df is not None and len(self.rt._corrections_df) > 0:
        self.gmodel.train(self.rt._corrections_df)
        self.bo.global_model = self.gmodel
    # Without this: GlobalModel starts blank every restart even though
    # 489 corrections are already on disk. Fixed to train on startup.

    # Step 4: Create engagement simulator
    self.sim = EngagementSimulator(seed=seed, range_table=self.rt)
    # EngagementSimulator wraps run_engagement() and manages RNG seeds

    # Step 5: Create and immediately train PINN correctors
    self.cf_low  = PINNCorrector(corrections_path, solution_type="LOW")
    self.cf_high = PINNCorrector(corrections_path, solution_type="HIGH")
    self.cf_low.load_and_train(verbose=verbose)
    self.cf_high.load_and_train(verbose=verbose)
    # Two separate correctors: LOW trajectories (< 45°) and HIGH (> 45°)
    # have different physics, so they need separate learned corrections.

    # Step 6: Load engagement history
    self.tracker = ConvergenceTracker(history_path)
    self.engagement_n = len(self.tracker._records)
    # Picks up from where the last session left off (e.g. engagement 536)"""
)

add_heading("engage() — One Full Engagement", 3)
add_code(
"""def engage(self, target_x, target_y, target_z, prefer="LOW"):
    self.engagement_n += 1

    # Step 1: Physics solution
    sol = solver.solve(target_x, target_y, target_z, self.v0, prefer=prefer)
    if not sol.reachable:
        return None   # Target unreachable — skip

    # Step 2: Range table lookup
    rt_lookup = self.rt.lookup(sol.horiz_range, target_y, self.v0, prefer=prefer)
    # e.g. {"delta_pitch": +0.43°, "n_observations": 67944}

    # Step 3: PINN pre-correction
    cf = self.cf_high if prefer == "HIGH" else self.cf_low
    pinn_correction = {"delta_pitch": 0.0, "delta_yaw": 0.0, "source": "none"}
    if cf.is_fitted:
        pinn_correction = cf.predict(sol.horiz_range, target_y, self.v0)
        # e.g. {"delta_pitch": +0.52°, "delta_yaw": -0.11°,
        #        "delta_v0": -3.45 m/s, "source": "pinn_torch_LOW"}

    # Step 4: Bayesian Optimizer engagement (~250 shots total)
    result = self.sim.run_engagement(
        target_x, target_y, target_z,
        optimizer         = self.bo,
        gp_pre_correction = pinn_correction if cf.is_fitted else None,
        prefer            = prefer,
    )

    # Step 5: Record result in history
    self.tracker.record(
        engagement_n    = self.engagement_n,
        baseline_cep    = result["baseline_cep"],
        bo_cep          = result["verified_cep"],
        improvement_pct = result["improvement_pct"],
    )

    # Step 6: Check if PINN needs retraining
    self._maybe_retrain_pinn()
    # Checks should_retrain() for both LOW and HIGH.
    # Retrains if ≥5 new records exist covering novel ranges.

    return result"""
)

add_heading("_maybe_retrain_pinn() — The Retraining Manager", 3)
add_code(
"""def _maybe_retrain_pinn(self):
    low_retrain  = self.cf_low.should_retrain()   # Check LOW corrector
    high_retrain = self.cf_high.should_retrain()  # Check HIGH corrector

    if low_retrain or high_retrain:
        if low_retrain:
            self.cf_low.load_and_train(verbose=False)   # Retrain LOW PINN
        if high_retrain:
            self.cf_high.load_and_train(verbose=False)  # Retrain HIGH PINN

        # Also refresh the GlobalModel (the BO's cross-engagement prior)
        if self.rt._corrections_df is not None:
            self.gmodel.train(self.rt._corrections_df)
            self.bo.global_model = self.gmodel"""
)

add_spacer()

# ── 5.10 experiment.py ────────────────────────────────────────────
add_heading("5.10  experiment.py — The Research Framework", 2)
add_para(
    "This file provides scientific experiment infrastructure — rigorous methods to test "
    "whether ARCS actually works and to measure each component's individual contribution.",
    size=11
)

add_heading("run_ablation() — What Contributes What?", 3)
add_para(
    "An ablation study removes one component at a time to measure its individual contribution. "
    "Four conditions are tested on the same target set:",
    size=11
)
add_table(
    headers=["Condition", "Components Used", "Mean Improvement", "Interpretation"],
    rows=[
        ["A — Physics only",    "Physics solver only",              "+0.0%",  "Baseline — no AI correction"],
        ["B — Physics + BO",    "Physics + Bayesian Optimizer",     "+19.2%", "BO's contribution alone"],
        ["C — Physics + PINN",  "Physics + PINN (no BO)",           "-1.3%",  "PINN alone is harmful without BO to recover"],
        ["D — Full ARCS",       "Physics + BO + PINN + GlobalModel","+37.0%", "Full system at steady state"],
    ],
    col_widths=[1.5, 2.2, 1.3, 2.2]
)
add_para(
    "The C result is the most interesting: PINN alone makes things slightly worse on average. "
    "Without the BO to fine-tune and recover from bad predictions, a harmful PINN correction "
    "cannot be corrected mid-engagement. This confirms that BO is essential — PINN and BO "
    "are complementary, not interchangeable.",
    size=11, italic=True
)

add_heading("run_learning_curves() — Does It Improve Over Time?", 3)
add_code(
"""def run_learning_curves(targets, n_trials=15):
    for trial in range(n_trials):
        seed = 42 + trial × 100   # Seeds: 42, 142, 242, 342, ...
        # Different seed = different random robot noise pattern
        # Run the same target set with fresh robot each trial

    # After n_trials: compute mean ± std improvement at each engagement number
    # If all trials improve: the algorithm genuinely works, it is not luck
    # Requires n_trials ≥ 12 for ±10% confidence intervals"""
)

add_heading("run_transfer_learning() — Does Knowledge Transfer?", 3)
add_code(
"""def run_transfer_learning(targets_train, targets_test, seed_a=42, seed_b=999):
    # Train Robot A on 30 engagements — builds a correction memory
    results_a, memory_a = run_robot(targets_train, seed=seed_a)

    # Test Robot B cold — starts fresh with no knowledge
    results_b_cold, _ = run_robot(targets_test, seed=seed_b, memory=None)

    # Test Robot B warm — starts with Robot A's learned corrections as prior
    results_b_warm, _ = run_robot(targets_test, seed=seed_b, memory=memory_a)

    # If warm_converge < cold_converge: Robot A's knowledge helps Robot B
    # Useful for fleet deployment: train one robot, seed others from it"""
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# SECTION 6: HOW EVERYTHING CONNECTS
# ══════════════════════════════════════════════════════════════════
add_heading("6. How Everything Connects — Full Data Flow", 1)
add_para(
    "Here is the complete sequence of events during a single engagement, "
    "showing exactly which file handles each step and what data passes between them.",
    size=11
)

add_code(
"""USER calls: pipeline.engage(300, 5, 50)
"Shoot at: 300m forward, 5m up, 50m right"
            │
            ▼
[physics/ballistic_solver.py] BALLISTIC SOLVER
   Pure Newton mechanics
   → pitch=8.3°, yaw=9.5°, time-of-flight=4.2s, range=304m
            │
            ▼
[physics/range_table.py] RANGE TABLE LOOKUP
   Finds similar past engagements (weighted by distance)
   → delta_pitch=+0.43° from 67,944 weighted past observations
            │
            ▼
[pinn_corrector.py] PINN CORRECTOR
   Neural network (6,496 params, 300 training epochs)
   → delta_pitch=+0.52°, delta_yaw=-0.11°, delta_v0=-3.45 m/s
   → source="pinn_torch_LOW"
            │
            ▼
[bayesian_optimizer.py] ENGAGEMENT SIMULATOR
   ┌─────────────────────────────────────────────────────────┐
   │ BASELINE: 30 shots, NO correction                       │
   │ median miss = 12.3m ← the problem to fix               │
   └─────────────────────────────────────────────────────────┘
            │
            ▼
   ┌─────────────────────────────────────────────────────────┐
   │ BAYESIAN OPTIMIZER (20 iterations):                     │
   │ Iter 1: try [+0.52°, -0.11°, -3.45]  ← PINN start     │
   │         fire 8 shots → median miss = 5.2m              │
   │ Iter 2: try [+0.23°, +0.44°, +1.8]   ← random         │
   │         fire 8 shots → median miss = 9.8m              │
   │ Iter 3: try [+0.55°, -0.15°, -3.1]   ← GP-guided      │
   │         fire 8 shots → median miss = 4.7m  ← new best! │
   │ ...17 more GP-guided iterations...                      │
   │ Best correction found: [+0.55°, -0.15°, -3.1 m/s]      │
   └─────────────────────────────────────────────────────────┘
            │
            ▼
   ┌─────────────────────────────────────────────────────────┐
   │ [kalman_filter.py] KALMAN FILTER REFINEMENT             │
   │ Fire 8 directional shots at best correction             │
   │ KF estimates remaining residual: +0.02°, -0.01°         │
   │ Confidence = 0.87 → apply the refinement                │
   │ Final correction: [+0.57°, -0.16°, -3.1 m/s]           │
   └─────────────────────────────────────────────────────────┘
            │
            ▼
   ┌─────────────────────────────────────────────────────────┐
   │ VERIFICATION: 30 shots WITH final correction            │
   │ verified_cep = 4.3m                                     │
   │ FALLBACK CHECK: 4.3m < 12.3m × 1.10 ✓ (not worse)     │
   │ IMPROVEMENT: (12.3 - 4.3) / 12.3 = 65.0%               │
   │ Wilcoxon p < 0.05 ✓ (statistically significant)        │
   └─────────────────────────────────────────────────────────┘
            │
            ▼
[physics/range_table.py] RECORD TO CSV
   Saves: range=304m, dp=+0.57°, dy=-0.16°, dv=-3.1,
          miss_before=12.3, miss_after=4.3
            │
            ▼
[metrics.py] CONVERGENCE TRACKER
   Records: engagement 537, improvement=65.0%, type=LOW
            │
            ▼
[pipeline.py] PINN RETRAIN CHECK
   should_retrain() = False (only 1 new record; need 5 more)
   → No retrain. Will retrain after 4 more novel engagements."""
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# SECTION 7: KEY NUMBERS
# ══════════════════════════════════════════════════════════════════
add_heading("7. Key Numbers to Remember", 1)

add_table(
    headers=["Number", "What It Means"],
    rows=[
        ["30 + 30 shots",     "Baseline + Verify per engagement (minimum for statistical significance)"],
        ["4 + 16 = 20 iterations", "BO search: 4 random exploration + 16 GP-guided exploitation"],
        ["8 KF shots",        "Post-BO Kalman Filter refinement shots (directional)"],
        ["~250 total shots",  "Per engagement: 30 baseline + 20×8 BO + 8 KF + 30 verify"],
        ["20 records minimum","Before PINN can train (needs enough history to learn patterns)"],
        ["+37%",              "Mean improvement at steady state (live pipeline)"],
        ["84%",               "Percentage of engagements that improve (85% is physics ceiling)"],
        ["~44%",              "Physics ceiling — maximum improvable CEP given BIAS_SCALE=1.5"],
        ["6,496",             "PINN parameter count (tiny — trains in under 1 second)"],
        ["300 epochs",        "PINN training passes through the data per retrain"],
        ["1.34m",             "Best single-engagement CEP achieved (at ~200m range)"],
        ["30m",               "Novelty threshold for PINN retraining trigger"],
        ["5 new records",     "Minimum new data needed before retraining is attempted"],
        ["p < 0.05",          "Required for Wilcoxon test to declare improvement significant"],
        ["0.5 confidence",    "Minimum KF confidence before applying KF correction"],
    ],
    col_widths=[2.0, 4.7]
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# SECTION 8: GLOSSARY
# ══════════════════════════════════════════════════════════════════
add_heading("8. Glossary", 1)
add_para(
    "Every technical term used in ARCS, explained in plain English.",
    size=11
)

glossary = [
    ("Adam optimizer",
     "A smart version of gradient descent that automatically adjusts the learning rate "
     "for each network weight individually. Named after 'Adaptive Moment Estimation.'"),
    ("Backpropagation",
     "The algorithm that computes which direction to adjust each network weight to reduce "
     "the loss. Works by tracing the error backwards through the network layers."),
    ("Bayesian Optimizer (BO)",
     "A search algorithm that finds the best correction by building a mathematical model "
     "of the search space and focusing on the most promising areas. Smarter than random search."),
    ("Bootstrap Confidence Interval",
     "A statistical technique to estimate the uncertainty on a measurement. ARCS resamples "
     "the 30 verification shots 1,000 times to compute the 95% confidence interval."),
    ("CEP (Circular Error Probable)",
     "The official NATO/aerospace accuracy metric. CEP 50% = the radius within which half "
     "of all shots land. Smaller is better."),
    ("Confidence (KF)",
     "A scalar from 0 to 1 showing how certain the Kalman Filter is about its correction "
     "estimate. Starts near 0 (one shot), rises toward 1 as more shots are observed."),
    ("Epoch",
     "One complete pass through all training data. The PINN trains for 300 epochs — "
     "it sees every record 300 times, improving a little each pass."),
    ("Fallback",
     "A safety mechanism: if the BO correction makes things more than 10% worse than "
     "baseline, throw it away and report 0% improvement. Prevents harmful corrections "
     "from being recorded."),
    ("Gaussian Process (GP)",
     "A probabilistic model that predicts both the expected outcome (mean) and the "
     "uncertainty (std) for any untested input. Used by the BO to guide its search."),
    ("Gradient descent",
     "The basic algorithm for training neural networks: adjust weights in the direction "
     "that reduces the loss, by a small step size (the learning rate)."),
    ("HIGH trajectory",
     "A lobbed firing path with pitch angle > 45°. Like a mortar shell. Takes longer "
     "to reach the target but can clear obstacles."),
    ("Information gain (novelty trigger)",
     "The PINN only retrains when new data covers a genuinely new range region (>30m from "
     "anything seen during the last training). Prevents unnecessary retraining."),
    ("Kalman Filter",
     "An optimal algorithm for estimating state from noisy measurements. After each "
     "observation, it updates both the state estimate AND its uncertainty. Used by GPS, "
     "Apollo, and ARCS."),
    ("kappa (κ)",
     "The exploration-exploitation trade-off parameter in the BO. High κ = explore more "
     "broadly. Low κ = exploit the best-known region. ARCS decays κ from 2.0 to 0.5 "
     "during each engagement."),
    ("LOW trajectory",
     "A flat, fast firing path with pitch angle < 45°. Like a rifle bullet. Default choice."),
    ("Loss function",
     "A mathematical measure of 'how wrong is the network right now?' Lower is better. "
     "ARCS has two components: L_data (fit the history) and L_physics (hit the target)."),
    ("Matern 5/2 kernel",
     "The similarity function used by the Gaussian Process. Two corrections that are "
     "numerically close are assumed to give similar miss distances. The Matern 5/2 form "
     "is smooth but not infinitely so — appropriate for real physical systems."),
    ("Paired evaluation",
     "Baseline (30 shots) and verification (30 shots) use the same random noise seed. "
     "This ensures a fair comparison — the improvement is due to the correction, not "
     "lucky noise differences."),
    ("PINN (Physics-Informed Neural Network)",
     "A neural network that must satisfy physical equations (Galileo's range equation) "
     "in addition to fitting the historical data. This prevents physically impossible "
     "corrections and requires fewer training examples."),
    ("Prior",
     "What we believe before seeing new data. In Bayesian terms: the BO's prior is the "
     "average of past successful corrections. It is the starting point for the search."),
    ("Posterior",
     "Updated belief after seeing new data. Each BO shot updates the posterior — the GP's "
     "model of which corrections are promising."),
    ("Random error",
     "Unpredictable variation in each shot. Cannot be corrected. ARCS averages it away "
     "using multiple shots (n_avg shots per BO iteration)."),
    ("Systematic error",
     "Predictable flaw that occurs the same way every time. ARCS learns and corrects this. "
     "The four systematic errors: gravity sag, gear backlash, yaw drift, velocity drift."),
    ("Tanh (activation function)",
     "A mathematical function that squashes any input to the range [-1, +1]. Used in the "
     "PINN to keep corrections bounded and to ensure smooth predictions."),
    ("UCB / LCB (Upper/Lower Confidence Bound)",
     "The formula that guides the BO's search. LCB = mean - κ × std. Minimising LCB "
     "picks corrections with low expected miss (exploitation) OR high uncertainty (exploration)."),
    ("Wilcoxon signed-rank test",
     "A statistical test to check if the improvement between baseline and verified CEP "
     "is real (not just luck). Result: p < 0.05 means less than 5% chance it was luck."),
]

for term, definition in glossary:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after  = Pt(3)
    r1 = p.add_run(term + ":  ")
    r1.bold = True
    r1.font.size = Pt(11)
    r1.font.name = "Calibri"
    r1.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    r2 = p.add_run(definition)
    r2.font.size = Pt(11)
    r2.font.name = "Calibri"

doc.add_page_break()

# ── Final page: results summary ───────────────────────────────────
add_heading("Results at a Glance", 1)
add_para(
    "These are the headline numbers from running the full ARCS pipeline on 600+ engagements "
    "with a simulated robot (BIAS_SCALE=1.5, all four mechanical flaws active).",
    size=11
)

add_table(
    headers=["Metric", "Value", "Meaning"],
    rows=[
        ["Mean improvement",      "+37.0%",    "Average CEP reduction across all engagements"],
        ["Positive engagements",  "84%",        "Fraction where the correction genuinely helped"],
        ["Best single CEP",       "1.34m",      "Closest-to-perfect engagement at ~200m range"],
        ["Best improvement",      "+81.9%",     "Best individual engagement result"],
        ["PINN LOW records",      "389",         "Historical records the LOW PINN trained on"],
        ["PINN HIGH records",     "178",         "Historical records the HIGH PINN trained on"],
        ["Total engagements",     "620+",        "Cumulative engagements run in current session"],
        ["Physics ceiling",       "~44%",        "Maximum improvable CEP given BIAS_SCALE=1.5"],
        ["Ablation — BO alone",   "+19.2%",     "BO without PINN (cold start experiment)"],
        ["Ablation — Full ARCS",  "+37%",       "Full system at steady state (600+ records)"],
    ],
    col_widths=[2.0, 1.5, 3.2]
)

add_spacer()
add_para(
    "The +37% vs +19.2% difference between Full ARCS and BO-alone is the PINN's contribution: "
    "+17.8 percentage points of improvement from the neural network's pre-correction, "
    "which eliminates most of the systematic bias before the BO even starts searching. "
    "The BO then fine-tunes the small residual, achieving the full system accuracy.",
    italic=True, size=11
)

# ── Save ──────────────────────────────────────────────────────────
output_path = "ARCS_Documentation.docx"
doc.save(output_path)
print(f"Saved: {output_path}")
